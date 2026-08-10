from io import BytesIO
from pathlib import Path

from docx import Document

from ..models import NormalizedUnit, ParsedDocument
from .base import DocumentParser


class DocxParser(DocumentParser):
    def parse(
        self,
        data: bytes,
        filename: str,
        mime_type: str | None,
    ) -> ParsedDocument:
        document = Document(
            BytesIO(data),
        )

        units: list[NormalizedUnit] = []

        section_index = 1
        current_lines: list[str] = []
        current_label = "Document"

        def flush_section() -> None:
            nonlocal section_index
            nonlocal current_lines
            nonlocal current_label

            text = "\n".join(
                current_lines,
            ).strip()

            if not text:
                return

            units.append(
                NormalizedUnit(
                    index=section_index,
                    kind="section",
                    label=current_label,
                    text=text,
                )
            )

            section_index += 1
            current_lines = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if not text:
                continue

            style_name = (
                paragraph.style.name
                if paragraph.style
                else ""
            )

            if style_name.lower().startswith(
                "heading"
            ):
                flush_section()
                current_label = text
                continue

            current_lines.append(text)

        flush_section()

        if not units:
            units.append(
                NormalizedUnit(
                    index=1,
                    kind="section",
                    label="Document",
                    text="",
                )
            )

        full_text = "\n\n".join(
            unit.text
            for unit in units
            if unit.text
        )

        return ParsedDocument(
            filename=filename,
            extension=Path(filename).suffix.lower(),
            mime_type=mime_type,
            parser="python-docx",
            units=units,
            full_text=full_text,
            metadata={
                "section_count": len(units),
            },
        )